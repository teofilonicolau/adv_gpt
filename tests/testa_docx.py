from docx import Document
import os

doc = Document()
doc.add_heading("Petição Teste", 0)
doc.add_paragraph("Esta é uma petição gerada para teste de abertura.")
doc.add_paragraph("Caso esteja lendo isso no Word, está tudo funcionando!")

# Caminho da pasta
pasta = "dados/peticoes_geradas"
os.makedirs(pasta, exist_ok=True)

# Salvar arquivo
caminho = os.path.join(pasta, "teste_peticao.docx")
doc.save(caminho)

print(f"✅ Arquivo gerado com sucesso: {caminho}")
