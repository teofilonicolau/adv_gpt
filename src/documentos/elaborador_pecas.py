from docx import Document
import datetime
import os

def gerar_peticao(nome_arquivo, titulo, qualificacao, fatos, fundamentos, pedidos, local="Icó/CE"):
    doc = Document()

    # Estrutura do documento
    doc.add_heading(titulo.upper(), level=1)
    doc.add_paragraph(qualificacao)
    doc.add_heading("I – DOS FATOS", level=2)
    doc.add_paragraph(fatos)
    doc.add_heading("II – DO DIREITO", level=2)
    doc.add_paragraph(fundamentos)
    doc.add_heading("III – DOS PEDIDOS", level=2)
    for p in pedidos:
        doc.add_paragraph(f"- {p}", style='List Bullet')
    data = datetime.datetime.now().strftime("%d de %B de %Y")
    doc.add_paragraph(f"\n\n{local}, {data}.")
    doc.add_paragraph("\n\n\nPrevInfoBot • Petição gerada automaticamente por IA jurídica\n", style='Intense Quote')

    # Salvar .docx
    pasta = os.path.join("dados", "peticoes_geradas")
    os.makedirs(pasta, exist_ok=True)
    caminho_docx = os.path.join(pasta, nome_arquivo)
    doc.save(caminho_docx)

    # PDF: Word se disponível, senão fallback
    metodo = None
    try:
        import shutil
        if shutil.which("winword"):
            from docx2pdf import convert
            convert(caminho_docx)
            metodo = "word"
            print("✅ PDF via Word gerado com sucesso")
        else:
            from src.documentos.docx2pdf_fallback import converter_pdf_sem_word
            converter_pdf_sem_word(caminho_docx)
            metodo = "fallback"
            print("✅ PDF via fallback gerado com sucesso")
    except Exception as e:
        metodo = "erro"
        print(f"⚠️ Erro ao gerar PDF: {e}")

    return caminho_docx, metodo
